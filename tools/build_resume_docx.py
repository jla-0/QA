from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\WritingPortfolio\QA_Manual_Tester\resume\Justin_Arndt_QA_Manual_Tester_Healthcare_Claims_ATS.docx"


def set_cell_text_safe(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.bold = bold
    return run


def set_paragraph_border(paragraph, color="2D5B9A"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(45, 91, 154)
    set_paragraph_border(p)
    return p


def add_body(doc, text, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.04
    run = p.add_run(text)
    run.font.size = Pt(9.3)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"- {text}")
    run.font.size = Pt(9.1)
    return p


def add_role(doc, title, company, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(9.8)
    r.font.color.rgb = RGBColor(16, 42, 67)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(f"{company} | {dates}")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(67, 80, 95)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor(16, 42, 67)
    normal.paragraph_format.space_after = Pt(3)

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Arial"
    list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    list_bullet.font.size = Pt(9.1)
    list_bullet.paragraph_format.left_indent = Inches(0.22)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.12)
    list_bullet.paragraph_format.space_after = Pt(1.5)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    r = name.add_run("Justin Arndt")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(16, 42, 67)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(1)
    r = contact.add_run("Lancaster, PA | 717-808-3119 | justinarndt05@gmail.com | LinkedIn: [add URL] | GitHub: [add URL] | Portfolio: [add URL]")
    r.font.size = Pt(8.8)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline.paragraph_format.space_after = Pt(5)
    r = headline.add_run("QA Manual Tester | Healthcare Claims | SQL Backend Testing | Regulated Systems QA")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(45, 91, 154)

    add_heading(doc, "Professional Summary")
    add_body(
        doc,
        "Software QA and regulated systems professional with 8+ years of validation, quality, systems testing, and technical documentation experience in FDA-regulated pharmaceutical environments, plus 10 years supporting healthcare IT systems in a HIPAA-aware clinical setting. Experienced translating business requirements and technical specifications into test plans, test cases, traceability matrices, SQL-backed evidence, defect documentation, and retest packages. Strong fit for healthcare claims QA roles requiring manual functional testing, backend database validation, regression testing, integration testing, XML/SOAP awareness, Jira-style defect tracking, and disciplined handling of sensitive healthcare data.",
    )

    add_heading(doc, "Core QA Strengths")
    strengths = (
        "Manual software QA and SQA testing | Healthcare claims QA portfolio proof | Business requirements analysis | Technical specification review | Test strategy, test plans, and test cases | Requirements traceability matrices | Functional, regression, and integration testing | Web and GUI application testing | Backend database testing and SQL validation | Oracle SQL and SQL Server query concepts | Defect documentation, triage, retest, and closure | Jira-style defect tracking and evidence packages | XML, SOAP, SOAP UI, and web services testing concepts | Healthcare EDI awareness including 837, 835, 270/271, and 276/277 flows | HIPAA-aware documentation, PHI safeguards, and audit-ready evidence"
    )
    add_body(doc, strengths)

    add_heading(doc, "Technical Skills")
    add_body(doc, "Testing and QA: Manual testing, SQA, functional testing, regression testing, integration testing, backend testing, requirements-based testing, negative testing, boundary testing, user acceptance support, defect lifecycle management, retesting, validation evidence, test summaries")
    add_body(doc, "Data and SQL: SQL, complex joins, CTE concepts, aggregations, reconciliation queries, data integrity checks, ETL concepts, data migration support, SQL-based reporting, Oracle SQL concepts, SQL Server/T-SQL concepts")
    add_body(doc, "Healthcare and Compliance: Healthcare IT, HIPAA-aware workflows, PHI handling discipline, PACS, DICOM, clinical systems, medical terminology exposure, regulated documentation, audit trails, access controls, electronic signatures, data integrity")
    add_body(doc, "Tools and Technologies: Jira concepts, Git, GitHub, Markdown, XML, SOAP, SOAP UI concepts, REST concepts, Python, Power BI, Bash, R, AWS, Azure, Docker, Veeva Vault, Documentum, LIMS, chromatography data systems, environmental monitoring systems")

    add_heading(doc, "Professional Experience")
    add_role(
        doc,
        "Senior Technical Writer and Computer System Validation Engineer / QA Engineer",
        "GSK, contract via Alphanumeric Systems, Inc. | Marietta, PA",
        "Jan 2016 - Apr 2024",
    )
    for item in [
        "Planned, authored, and executed requirements-based validation and SQA documentation for 10+ mission-critical laboratory and quality systems used in regulated manufacturing operations.",
        "Created user requirements, functional specifications, test protocols, test cases, traceability matrices, IQ/OQ/PQ evidence, validation summaries, SOPs, change controls, deviation records, and defect-style investigation writeups.",
        "Performed functional, regression, integration, configuration, access control, audit trail, electronic signature, backup/restore, upgrade, and data integrity testing across standalone and client-server systems.",
        "Used SQL, Python, and structured data review to validate backend behavior, reconcile migrated records, investigate data anomalies, and support evidence-based quality decisions.",
        "Reverse-engineered legacy SQL and VB-based tools to document system behavior, validate data flows, and support migration and modernization planning.",
        "Documented defects, deviations, impact assessments, root cause, corrective actions, retest evidence, and effectiveness checks for QA and regulatory review.",
        "Partnered with QA, IT, scientists, vendors, and business stakeholders to resolve issues, retest fixes, protect validated-state operation, and communicate quality risk clearly.",
        "Standardized templates and documentation frameworks that reduced authoring time by 30% while improving consistency, review quality, and reuse across cross-functional teams.",
        "Developed and maintained 150+ SOPs and controlled procedures covering administration, backup and restore, audit trail review, user access management, data integrity verification, and system support.",
        "Supported migrations, software upgrades, and patching efforts under formal change control with zero unplanned downtime and clear post-change verification evidence.",
        "Supported multiple audits with zero documentation-related findings during tenure.",
        "Mentored junior engineers and cross-functional contributors on validation methodology, test evidence quality, technical documentation standards, and compliance expectations.",
    ]:
        add_bullet(doc, item)

    add_role(
        doc,
        "Technical Documentation Coordinator",
        "Lancaster General Health / Penn Medicine | Lancaster, PA",
        "Apr 2004 - Jun 2014",
    )
    for item in [
        "Created technical procedures, user guides, support documentation, troubleshooting workflows, and continuity guidance for PACS, DICOM imaging systems, and healthcare IT workflows supporting 100+ clinical and technical users.",
        "Documented system integrations, operational procedures, backup and recovery steps, access-related processes, and support workflows in a HIPAA-aware healthcare environment.",
        "Collaborated with IT, clinical users, and support teams to clarify issues, document repeatable resolution paths, and improve system support consistency.",
        "Built onboarding and standards documentation that reduced training time by 30%.",
        "Created self-service support content that reduced support ticket volume by 40%.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Relevant Portfolio Project")
    add_role(
        doc,
        "Healthcare Claims Manual QA Portfolio",
        "Independent portfolio project",
        "2026",
    )
    for item in [
        "Built a GitHub-ready manual QA project that demonstrates how to test a synthetic healthcare claims workflow from requirement review through release recommendation.",
        "Created a healthcare claims test strategy, test plan, requirements traceability matrix, manual test cases, SQL backend validation checks, Jira-style defect examples, HIPAA-safe test data strategy, EDI samples, and SOAP/XML service test examples.",
        "Covered claim intake, eligibility, provider contract validation, adjudication, denial handling, claim status inquiry, remittance advice, regression risk, integration points, and load testing considerations.",
        "Designed the project with synthetic data only, showing PHI discipline and safe public portfolio presentation for a high-risk healthcare data role.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Education")
    add_body(doc, "Master of Science, Data Analytics / Data Science | Western Governors University | 2025")
    add_body(doc, "Bachelor of Science, Cybersecurity and Information Assurance | Western Governors University | 2024")

    add_heading(doc, "ATS Keyword Set")
    add_body(
        doc,
        "QA Manual Tester, Software Quality Assurance, SQA, healthcare claims, medical claims, claims processing, claim adjudication, backend testing, database testing, complex SQL, Oracle SQL, SQL Server, regression testing, integration testing, functional testing, Web testing, GUI testing, XML, SOAP, SOAP UI, web services testing, EDI, 837, 835, 270/271, 276/277, Jira, defect tracking, defect management, test strategy, test plan, test cases, traceability matrix, HIPAA, PHI, medical terminology, load testing, performance testing, quality assurance, root cause analysis, data analysis, production support, validation, regulated systems, audit trails, access controls, change control.",
        after=0,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
